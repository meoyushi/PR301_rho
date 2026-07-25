"""Regenerate tests/fixtures/simple.docx. Run: python tests/fixtures/make_simple_docx.py"""

from pathlib import Path

from docx import Document

OUT = Path(__file__).parent / "simple.docx"


def build() -> None:
    doc = Document()
    doc.add_heading("Alice Johnson", level=1)
    doc.add_paragraph("Senior Python Engineer")
    doc.add_paragraph("Skills: Python, FastAPI, AWS")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Acme Corp - Backend Engineer 2019-2022")
    doc.add_paragraph("Built REST APIs serving 2M requests per day")
    doc.add_paragraph("Globex Inc - Software Engineer 2017-2019")
    doc.add_paragraph("Built REST APIs serving 2M requests per day")
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(f"wrote {OUT}")
