"""Render a `StructuredResume` to a .docx byte stream (python-docx).

Deliberately plain: the export is a clean, ATS-parseable Word document — one
section per résumé section, no columns, no text boxes, no tables (those are the
constructs that break naive ATS parsers). Styling is intentionally minimal and
does NOT mirror the browser preview pixel-for-pixel; the goal is a recruiter-
and-parser-friendly file, not a screenshot.

`section_order` mirrors the editor's style setting so the Word file lays sections
out in the same order the user arranged on screen.
"""

import io

from docx import Document
from docx.shared import Pt, RGBColor

from rho.models.resume import StructuredResume

_DEFAULT_ORDER = ["summary", "skills", "work", "projects", "achievements", "education"]


def _accent(hex_str: str) -> RGBColor:
    h = (hex_str or "#1c1b19").lstrip("#")
    if len(h) != 6:
        h = "1c1b19"
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _heading(doc: Document, text: str, color: RGBColor) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = color


def _bullet(doc: Document, text: str) -> None:
    if not text or not text.strip():
        return
    doc.add_paragraph(text.strip(), style="List Bullet")


def build_docx(resume: StructuredResume, section_order: list[str] | None = None,
               accent: str = "#b5482a", hidden_sections: list[str] | None = None) -> bytes:
    """Return the résumé as .docx bytes.

    `hidden_sections` are dropped from the output entirely — including from the
    fallback that would otherwise re-append any section missing from
    `section_order`. This lets the editor's show/hide toggles take effect in the
    Word file without deleting the underlying résumé data.
    """
    order = section_order or _DEFAULT_ORDER
    hide = set(hidden_sections or [])
    color = _accent(accent)
    doc = Document()

    # Name + headline header.
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(resume.name or "")
    name_run.bold = True
    name_run.font.size = Pt(20)
    if resume.headline:
        hp = doc.add_paragraph()
        hr = hp.add_run(resume.headline)
        hr.font.size = Pt(11)
        hr.font.color.rgb = color

    contact = [*(resume.emails or []), *(resume.phones or []), *(resume.urls or [])]
    if contact:
        cp = doc.add_paragraph()
        cp.add_run(" · ".join(contact)).font.size = Pt(9)

    def render(section: str) -> None:
        if section == "summary" and resume.summary:
            _heading(doc, "Summary", color)
            doc.add_paragraph(resume.summary.strip())
        elif section == "skills" and resume.skills:
            _heading(doc, "Skills", color)
            doc.add_paragraph(", ".join(resume.skills))
        elif section == "work" and resume.work:
            _heading(doc, "Experience", color)
            for w in resume.work:
                line = doc.add_paragraph()
                r = line.add_run(w.title or "")
                r.bold = True
                dates = " – ".join(x for x in [w.start_date, w.end_date] if x)
                meta = w.company + (f"   {dates}" if dates else "") if w.company else dates
                if meta:
                    sub = doc.add_paragraph()
                    sr = sub.add_run(meta)
                    sr.italic = True
                    sr.font.size = Pt(10)
                for b in w.bullets or []:
                    _bullet(doc, b)
        elif section == "projects" and resume.projects:
            _heading(doc, "Projects", color)
            for p in resume.projects:
                line = doc.add_paragraph()
                line.add_run(p.name or "").bold = True
                if p.url:
                    line.add_run(f"   {p.url}").font.size = Pt(9)
                if p.tech:
                    tp = doc.add_paragraph()
                    tr = tp.add_run(", ".join(p.tech))
                    tr.italic = True
                    tr.font.size = Pt(10)
                for b in p.bullets or []:
                    _bullet(doc, b)
        elif section == "achievements" and resume.achievements:
            _heading(doc, "Achievements", color)
            for a in resume.achievements:
                _bullet(doc, a)
        elif section == "education" and resume.education:
            _heading(doc, "Education", color)
            for e in resume.education:
                parts = [e.institution]
                if e.degree:
                    parts.append(e.degree)
                if e.field:
                    parts.append(e.field)
                tail = ", ".join(p for p in parts if p)
                if e.end_year:
                    tail += f" ({e.end_year})"
                doc.add_paragraph(tail)

    for section in order:
        if section not in hide:
            render(section)
    # Render any section the order list omitted, so nothing is silently lost —
    # unless it was explicitly hidden via the editor's toggle.
    for section in _DEFAULT_ORDER:
        if section not in order and section not in hide:
            render(section)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
